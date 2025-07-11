"""
知识提取器 - 从文档中提取知识点和构建知识图谱
"""

import os
import re
import json
import hashlib
from typing import List, Dict, Set, Tuple, Optional, Any
from pathlib import Path
from dataclasses import dataclass, asdict
import logging
from collections import defaultdict

import PyPDF2
from docx import Document
import markdown
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)

@dataclass
class ExtractedKnowledgePoint:
    """提取的知识点数据结构"""
    name: str
    description: str
    content: str
    difficulty_level: int
    keywords: List[str]
    learning_objectives: List[str]
    examples: str
    chapter: str
    section: str
    references: List[str]
    dependencies: List[str]  # 依赖的其他知识点名称

@dataclass
class ExtractedDomain:
    """提取的知识域数据结构"""
    name: str
    description: str
    knowledge_points: List[ExtractedKnowledgePoint]
    exam_weight: float
    sort_order: int

class DocumentParser:
    """文档解析器基类"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析文档，返回结构化数据"""
        raise NotImplementedError

class PDFParser(DocumentParser):
    """PDF文档解析器"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                content = ""
                for page in reader.pages:
                    content += page.extract_text() + "\n"
                
                return {
                    'content': content,
                    'type': 'pdf',
                    'pages': len(reader.pages),
                    'metadata': reader.metadata if hasattr(reader, 'metadata') else {}
                }
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            return {'content': '', 'type': 'pdf', 'error': str(e)}

class DocxParser(DocumentParser):
    """Word文档解析器"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        try:
            doc = Document(file_path)
            content = ""
            structure = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    content += text + "\n"
                    
                    # 尝试识别标题层级
                    if para.style.name.startswith('Heading'):
                        level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                        structure.append({
                            'text': text,
                            'level': level,
                            'type': 'heading'
                        })
                    else:
                        structure.append({
                            'text': text,
                            'type': 'paragraph'
                        })
            
            return {
                'content': content,
                'type': 'docx',
                'structure': structure,
                'paragraphs': len(doc.paragraphs)
            }
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            return {'content': '', 'type': 'docx', 'error': str(e)}

class MarkdownParser(DocumentParser):
    """Markdown文档解析器"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
                # 转换为HTML以便提取结构
                html = markdown.markdown(content, extensions=['toc'])
                soup = BeautifulSoup(html, 'html.parser')
                
                structure = []
                for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                    level = int(heading.name[1])
                    structure.append({
                        'text': heading.get_text(),
                        'level': level,
                        'type': 'heading'
                    })
                
                return {
                    'content': content,
                    'type': 'markdown',
                    'structure': structure,
                    'html': html
                }
        except Exception as e:
            logger.error(f"Error parsing Markdown {file_path}: {e}")
            return {'content': '', 'type': 'markdown', 'error': str(e)}

class TextParser(DocumentParser):
    """纯文本解析器"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
                return {
                    'content': content,
                    'type': 'text',
                    'lines': len(content.split('\n'))
                }
        except Exception as e:
            logger.error(f"Error parsing text {file_path}: {e}")
            return {'content': '', 'type': 'text', 'error': str(e)}

class KnowledgeExtractor:
    """知识提取器主类"""
    
    def __init__(self):
        self.parsers = {
            '.pdf': PDFParser(),
            '.docx': DocxParser(),
            '.doc': DocxParser(),
            '.md': MarkdownParser(),
            '.markdown': MarkdownParser(),
            '.txt': TextParser(),
        }
        
        # 知识点模式匹配规则
        self.knowledge_patterns = {
            'chapter': [
                r'第\s*(\d+)\s*章\s*(.+)',
                r'Chapter\s*(\d+)\s*[:\-]?\s*(.+)',
                r'(\d+)\.\s*(.+?)(?=\n|$)',
            ],
            'section': [
                r'(\d+)\.(\d+)\s*(.+)',
                r'§\s*(\d+)\.(\d+)\s*(.+)',
                r'Section\s*(\d+)\.(\d+)\s*(.+)',
            ],
            'knowledge_point': [
                r'(?:理解|掌握|了解|熟悉)(.+?)(?=。|；|，|\n)',
                r'(?:学习目标|目标|要求)[:：]\s*(.+?)(?=\n\n|\n[A-Z]|\n\d)',
                r'(?:知识点|要点|重点)[:：]\s*(.+?)(?=\n\n|\n[A-Z]|\n\d)',
            ],
            'difficulty_keywords': {
                '了解': 1,
                '理解': 2,
                '掌握': 3,
                '熟练掌握': 4,
                '精通': 5,
                'understand': 2,
                'master': 3,
                'proficient': 4,
                'expert': 5
            },
            'dependency_keywords': [
                r'(?:需要|要求|必须|应当|应该)(?:先|首先|预先)(?:学习|掌握|了解)(.+?)(?=。|；|，|\n)',
                r'(?:基于|基础|前提|前置)(?:是|为)(.+?)(?=。|；|，|\n)',
                r'(?:prerequisite|require|need|based on)(?:\s+)(.+?)(?=\.|;|,|\n)',
            ]
        }
    
    def extract_from_directory(self, directory_path: str, 
                             progress_callback: Optional[callable] = None) -> List[ExtractedDomain]:
        """从目录中提取知识体系"""
        results = []
        files = list(self._scan_files(directory_path))
        total_files = len(files)
        
        logger.info(f"开始提取知识体系，共发现 {total_files} 个文档文件")
        
        for i, file_path in enumerate(files):
            try:
                if progress_callback:
                    progress_callback(i / total_files * 0.8, f"处理文件: {os.path.basename(file_path)}")
                
                domain = self._extract_from_file(file_path)
                if domain and domain.knowledge_points:
                    results.append(domain)
                    logger.info(f"从 {file_path} 提取到 {len(domain.knowledge_points)} 个知识点")
                
            except Exception as e:
                logger.error(f"处理文件 {file_path} 时出错: {e}")
                continue
        
        # 后处理：构建依赖关系
        if progress_callback:
            progress_callback(0.9, "构建知识点依赖关系...")
        
        self._build_dependencies(results)
        
        if progress_callback:
            progress_callback(1.0, f"完成！提取到 {len(results)} 个知识域")
        
        return results
    
    def _scan_files(self, directory_path: str) -> List[str]:
        """扫描目录中的文档文件"""
        supported_extensions = set(self.parsers.keys())
        files = []
        
        for root, dirs, filenames in os.walk(directory_path):
            for filename in filenames:
                if any(filename.lower().endswith(ext) for ext in supported_extensions):
                    files.append(os.path.join(root, filename))
        
        return files
    
    def _extract_from_file(self, file_path: str) -> Optional[ExtractedDomain]:
        """从单个文件提取知识域"""
        file_ext = Path(file_path).suffix.lower()
        parser = self.parsers.get(file_ext)
        
        if not parser:
            logger.warning(f"不支持的文件格式: {file_ext}")
            return None
        
        # 解析文档
        parsed_data = parser.parse(file_path)
        if 'error' in parsed_data:
            logger.error(f"解析文件失败: {parsed_data['error']}")
            return None
        
        content = parsed_data['content']
        if not content.strip():
            logger.warning(f"文件内容为空: {file_path}")
            return None
        
        # 提取知识域信息
        domain_name = self._extract_domain_name(file_path, content)
        domain_description = self._extract_domain_description(content)
        
        # 提取知识点
        knowledge_points = self._extract_knowledge_points(content, parsed_data)
        
        if not knowledge_points:
            logger.warning(f"未能从文件中提取到知识点: {file_path}")
            return None
        
        return ExtractedDomain(
            name=domain_name,
            description=domain_description,
            knowledge_points=knowledge_points,
            exam_weight=self._calculate_exam_weight(content),
            sort_order=self._calculate_sort_order(file_path)
        )
    
    def _extract_domain_name(self, file_path: str, content: str) -> str:
        """提取知识域名称"""
        # 首先尝试从文件名提取
        filename = Path(file_path).stem
        
        # 清理文件名
        domain_name = re.sub(r'[\d\-_]', ' ', filename).strip()
        
        # 尝试从内容中提取更好的标题
        title_patterns = [
            r'第\s*(\d+)\s*章\s*(.+?)(?=\n|第\d+章)',
            r'Chapter\s*(\d+)\s*[:\-]?\s*(.+?)(?=\n|Chapter)',
            r'^#\s*(.+?)(?=\n|$)',  # Markdown一级标题
            r'^(.+?)(?=\n=+\n)',    # 文本下划线标题
        ]
        
        for pattern in title_patterns:
            match = re.search(pattern, content, re.MULTILINE | re.IGNORECASE)
            if match:
                if len(match.groups()) > 1:
                    domain_name = match.group(2).strip()
                else:
                    domain_name = match.group(1).strip()
                break
        
        return domain_name[:100] if domain_name else "未命名知识域"
    
    def _extract_domain_description(self, content: str) -> str:
        """提取知识域描述"""
        # 尝试提取前几段作为描述
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        if len(paragraphs) > 1:
            # 跳过标题，取前几段
            description_parts = []
            for para in paragraphs[1:4]:  # 最多取3段
                if len(para) > 20 and len(para) < 500:  # 合理长度的段落
                    description_parts.append(para)
            
            return '\n'.join(description_parts)[:1000]
        
        return "从文档自动提取的知识域"
    
    def _extract_knowledge_points(self, content: str, parsed_data: Dict) -> List[ExtractedKnowledgePoint]:
        """提取知识点"""
        knowledge_points = []
        
        # 基于文档结构提取
        if 'structure' in parsed_data:
            knowledge_points.extend(self._extract_from_structure(parsed_data['structure'], content))
        
        # 基于模式匹配提取
        knowledge_points.extend(self._extract_from_patterns(content))
        
        # 去重和清理
        knowledge_points = self._deduplicate_knowledge_points(knowledge_points)
        
        return knowledge_points
    
    def _extract_from_structure(self, structure: List[Dict], content: str) -> List[ExtractedKnowledgePoint]:
        """基于文档结构提取知识点"""
        knowledge_points = []
        current_chapter = ""
        current_section = ""
        
        for i, item in enumerate(structure):
            if item['type'] == 'heading':
                level = item.get('level', 1)
                text = item['text']
                
                if level <= 2:  # 章节级别
                    current_chapter = text
                    current_section = ""
                elif level <= 4:  # 小节级别
                    current_section = text
                
                # 从标题中提取可能的知识点
                kp = self._create_knowledge_point_from_heading(
                    text, level, current_chapter, current_section, content
                )
                if kp:
                    knowledge_points.append(kp)
        
        return knowledge_points
    
    def _extract_from_patterns(self, content: str) -> List[ExtractedKnowledgePoint]:
        """基于模式匹配提取知识点"""
        knowledge_points = []
        
        # 按段落分割内容
        paragraphs = content.split('\n\n')
        current_chapter = ""
        current_section = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # 识别章节
            for pattern in self.knowledge_patterns['chapter']:
                match = re.search(pattern, para)
                if match:
                    current_chapter = match.group(2) if len(match.groups()) > 1 else match.group(1)
                    break
            
            # 识别小节
            for pattern in self.knowledge_patterns['section']:
                match = re.search(pattern, para)
                if match:
                    current_section = match.group(3) if len(match.groups()) > 2 else match.group(2)
                    break
            
            # 识别知识点
            for pattern in self.knowledge_patterns['knowledge_point']:
                matches = re.finditer(pattern, para)
                for match in matches:
                    name = match.group(1).strip()
                    if len(name) > 5 and len(name) < 200:  # 合理长度
                        kp = self._create_knowledge_point(
                            name, para, current_chapter, current_section, content
                        )
                        knowledge_points.append(kp)
        
        return knowledge_points
    
    def _create_knowledge_point_from_heading(self, heading: str, level: int, 
                                           chapter: str, section: str, content: str) -> Optional[ExtractedKnowledgePoint]:
        """从标题创建知识点"""
        # 过滤掉过短或过长的标题
        if len(heading) < 3 or len(heading) > 200:
            return None
        
        # 跳过明显的章节标题
        if level <= 2 and any(keyword in heading for keyword in ['章', 'Chapter', '部分', 'Part']):
            return None
        
        return self._create_knowledge_point(heading, "", chapter, section, content)
    
    def _create_knowledge_point(self, name: str, description: str, 
                              chapter: str, section: str, content: str) -> ExtractedKnowledgePoint:
        """创建知识点对象"""
        # 提取关键词
        keywords = self._extract_keywords(name + " " + description)
        
        # 确定难度等级
        difficulty = self._determine_difficulty(name + " " + description)
        
        # 提取学习目标
        objectives = self._extract_learning_objectives(description, content)
        
        # 提取示例
        examples = self._extract_examples(description, content)
        
        # 提取参考资料
        references = self._extract_references(description, content)
        
        # 提取依赖关系
        dependencies = self._extract_dependencies(name + " " + description, content)
        
        return ExtractedKnowledgePoint(
            name=name.strip(),
            description=description.strip()[:1000],
            content=self._extract_detailed_content(name, content)[:2000],
            difficulty_level=difficulty,
            keywords=keywords,
            learning_objectives=objectives,
            examples=examples,
            chapter=chapter,
            section=section,
            references=references,
            dependencies=dependencies
        )
    
    def _extract_keywords(self, text: str) -> List[str]:
        """提取关键词"""
        # 简单的关键词提取，可以后续优化为更复杂的NLP方法
        keywords = []
        
        # 技术术语模式
        tech_patterns = [
            r'[A-Z]{2,}(?:[A-Z][a-z]+)*',  # 大写缩写
            r'[a-zA-Z]+-[a-zA-Z]+',       # 连字符术语
            r'[\u4e00-\u9fa5]{2,8}(?:技术|方法|模式|架构|设计|原理|算法)', # 中文技术术语
        ]
        
        for pattern in tech_patterns:
            matches = re.findall(pattern, text)
            keywords.extend(matches)
        
        # 去重并限制数量
        keywords = list(set(keywords))[:10]
        
        return keywords
    
    def _determine_difficulty(self, text: str) -> int:
        """确定难度等级"""
        difficulty = 2  # 默认难度
        
        for keyword, level in self.knowledge_patterns['difficulty_keywords'].items():
            if keyword in text:
                difficulty = max(difficulty, level)
        
        return min(difficulty, 5)
    
    def _extract_learning_objectives(self, description: str, content: str) -> List[str]:
        """提取学习目标"""
        objectives = []
        
        objective_patterns = [
            r'(?:学习目标|目标|要求)[:：]\s*(.+?)(?=\n\n|\n[A-Z]|\n\d)',
            r'(?:能够|可以|会)(.+?)(?=。|；|，|\n)',
            r'(?:understand|learn|master|know)(.+?)(?=\.|;|,|\n)',
        ]
        
        for pattern in objective_patterns:
            matches = re.finditer(pattern, description + " " + content, re.IGNORECASE)
            for match in matches:
                obj = match.group(1).strip()
                if 5 < len(obj) < 200:
                    objectives.append(obj)
        
        return objectives[:5]  # 最多5个目标
    
    def _extract_examples(self, description: str, content: str) -> str:
        """提取示例"""
        example_patterns = [
            r'(?:例如|示例|举例|比如)[:：]?\s*(.+?)(?=\n\n|\n[A-Z]|\n\d)',
            r'(?:example|for instance|such as)[::]?\s*(.+?)(?=\n\n|\n[A-Z]|\n\d)',
        ]
        
        examples = []
        for pattern in example_patterns:
            matches = re.finditer(pattern, description + " " + content, re.IGNORECASE)
            for match in matches:
                example = match.group(1).strip()
                if 10 < len(example) < 500:
                    examples.append(example)
        
        return '\n'.join(examples[:3])[:1000]  # 最多3个示例
    
    def _extract_references(self, description: str, content: str) -> List[str]:
        """提取参考资料"""
        ref_patterns = [
            r'(?:参考|参见|见|refer to|see|reference)[:：]?\s*(.+?)(?=\n|$)',
            r'\[(\d+)\]',  # 数字引用
            r'（(.+?)）',  # 括号引用
        ]
        
        references = []
        for pattern in ref_patterns:
            matches = re.finditer(pattern, description + " " + content)
            for match in matches:
                ref = match.group(1).strip()
                if 3 < len(ref) < 100:
                    references.append(ref)
        
        return list(set(references[:10]))  # 去重，最多10个
    
    def _extract_dependencies(self, text: str, content: str) -> List[str]:
        """提取依赖关系"""
        dependencies = []
        
        for pattern in self.knowledge_patterns['dependency_keywords']:
            matches = re.finditer(pattern, text + " " + content, re.IGNORECASE)
            for match in matches:
                dep = match.group(1).strip()
                if 3 < len(dep) < 100:
                    dependencies.append(dep)
        
        return list(set(dependencies[:5]))  # 去重，最多5个依赖
    
    def _extract_detailed_content(self, name: str, content: str) -> str:
        """提取详细内容"""
        # 在内容中查找与知识点名称相关的段落
        paragraphs = content.split('\n\n')
        relevant_paragraphs = []
        
        for para in paragraphs:
            # 简单的相关性检查
            if any(word in para for word in name.split() if len(word) > 2):
                relevant_paragraphs.append(para.strip())
        
        return '\n\n'.join(relevant_paragraphs[:3])[:2000]  # 最多3段
    
    def _deduplicate_knowledge_points(self, knowledge_points: List[ExtractedKnowledgePoint]) -> List[ExtractedKnowledgePoint]:
        """去重知识点"""
        seen_names = set()
        unique_points = []
        
        for kp in knowledge_points:
            # 标准化名称用于比较
            normalized_name = re.sub(r'\s+', ' ', kp.name.lower().strip())
            
            if normalized_name not in seen_names and len(normalized_name) > 3:
                seen_names.add(normalized_name)
                unique_points.append(kp)
        
        return unique_points
    
    def _calculate_exam_weight(self, content: str) -> float:
        """计算考试权重"""
        # 基于内容长度和重要性关键词计算权重
        weight = min(len(content) / 10000, 1.0)  # 基础权重基于内容长度
        
        # 重要性关键词加权
        important_keywords = ['重点', '核心', '关键', '重要', 'important', 'key', 'core', 'critical']
        for keyword in important_keywords:
            if keyword in content:
                weight += 0.1
        
        return min(weight, 1.0)
    
    def _calculate_sort_order(self, file_path: str) -> int:
        """计算排序顺序"""
        filename = os.path.basename(file_path)
        
        # 尝试从文件名中提取数字
        numbers = re.findall(r'\d+', filename)
        if numbers:
            return int(numbers[0])
        
        return 999  # 默认排序
    
    def _build_dependencies(self, domains: List[ExtractedDomain]):
        """构建知识点间的依赖关系"""
        # 创建所有知识点的名称映射
        all_knowledge_points = {}
        for domain in domains:
            for kp in domain.knowledge_points:
                all_knowledge_points[kp.name.lower()] = kp
        
        # 解析依赖关系
        for domain in domains:
            for kp in domain.knowledge_points:
                resolved_deps = []
                for dep_name in kp.dependencies:
                    # 尝试在所有知识点中找到匹配的依赖
                    for name, target_kp in all_knowledge_points.items():
                        if dep_name.lower() in name or name in dep_name.lower():
                            if target_kp.name != kp.name:  # 避免自依赖
                                resolved_deps.append(target_kp.name)
                            break
                
                kp.dependencies = list(set(resolved_deps))  # 去重
    
    def generate_knowledge_graph_data(self, domains: List[ExtractedDomain]) -> Dict[str, Any]:
        """生成知识图谱数据"""
        nodes = []
        edges = []
        
        # 添加知识域节点
        for domain in domains:
            nodes.append({
                'id': f"domain_{domain.name}",
                'name': domain.name,
                'type': 'domain',
                'description': domain.description,
                'exam_weight': domain.exam_weight,
                'knowledge_point_count': len(domain.knowledge_points)
            })
            
            # 添加知识点节点
            for kp in domain.knowledge_points:
                nodes.append({
                    'id': f"kp_{kp.name}",
                    'name': kp.name,
                    'type': 'knowledge_point',
                    'description': kp.description,
                    'difficulty_level': kp.difficulty_level,
                    'chapter': kp.chapter,
                    'section': kp.section,
                    'keywords': kp.keywords
                })
                
                # 添加域到知识点的边
                edges.append({
                    'source': f"domain_{domain.name}",
                    'target': f"kp_{kp.name}",
                    'type': 'contains'
                })
                
                # 添加知识点依赖边
                for dep in kp.dependencies:
                    edges.append({
                        'source': f"kp_{dep}",
                        'target': f"kp_{kp.name}",
                        'type': 'prerequisite'
                    })
        
        return {
            'nodes': nodes,
            'edges': edges,
            'statistics': {
                'domain_count': len(domains),
                'knowledge_point_count': sum(len(d.knowledge_points) for d in domains),
                'dependency_count': len([e for e in edges if e['type'] == 'prerequisite'])
            }
        } 